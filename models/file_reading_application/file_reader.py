from pathlib import Path
import os
import pdfplumber
from docx import Document

UNPROCESSED_DIR = Path(__file__).parent / "unprocessed_data"
PROCESSED_DIR = Path(__file__).parent / "processed_data"

def read_pdf(file_path):
    text_parts = []
    
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
            
            tables = page.extract_tables()
            for table in tables:
                for row in table:
                    row_text = "\t".join(cell or "" for cell in row)
                    text_parts.append(row_text)
    
    return "\n".join(text_parts)
    
def read_docx(file_path):
    doc = Document(file_path)
    text_parts = []
    
    for paragraph in doc.paragraphs:
        stripped_paragraph = paragraph.text.strip()
        if stripped_paragraph:
            text_parts.append(stripped_paragraph)
    
    for table in doc.tables:
        for row in table.rows:
            row_text = "\t".join(cell.text.strip() for cell in row.cells)
            text_parts.append(row_text)
    
    return "\n".join(text_parts)

def save_text(text, output_path):
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(text)

def main():
    for filename in os.listdir(UNPROCESSED_DIR):
        file_path = UNPROCESSED_DIR/filename
        if filename.lower().endswith(".pdf"):
            text = read_pdf(file_path)
        elif filename.lower().endswith(".docx"):
            text = read_docx(file_path)
        else:
            continue
        
        output_file = PROCESSED_DIR/(Path(filename).stem + ".txt")
        save_text(text, output_file)
        print(f"Processed: {filename}")

if __name__ == "__main__":
    main()